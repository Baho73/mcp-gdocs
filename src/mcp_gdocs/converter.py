# START_MODULE_CONTRACT
# MODULE_NAME: M-CONVERTER
# PURPOSE: Markdown → DOCX через python-docx + mistune
# INPUTS: markdown (str)
# OUTPUTS: docx_path (str) — path to temp DOCX file
# ERRORS: ConversionError
# END_MODULE_CONTRACT

"""Markdown to DOCX converter with full formatting support."""

import re
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


# START_BLOCK_INLINE_FORMATTING
def _apply_inline_formatting(paragraph, text: str):
    """Parse inline markdown (bold, italic, code, links) and add formatted runs."""
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\)|([^*`\[]+))"

    for match in re.finditer(pattern, text):
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        elif match.group(5):  # [text](url)
            run = paragraph.add_run(match.group(5))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        elif match.group(7):  # plain text
            paragraph.add_run(match.group(7))
# END_BLOCK_INLINE_FORMATTING


# START_BLOCK_CONVERT
def markdown_to_docx(markdown: str, output_path: str | None = None) -> str:
    """Convert Markdown string to a DOCX file.

    Args:
        markdown: Markdown-formatted text.
        output_path: Where to save. If None, creates a temp file.

    Returns:
        Path to the created DOCX file.
    """
    if output_path is None:
        fd, output_path = tempfile.mkstemp(suffix=".docx")
        import os
        os.close(fd)

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Heading styles
    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Horizontal rules
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Tables
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|", lines[i + 1]):
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header + separator

            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row_cells)
                i += 1

            num_cols = len(header_cells)
            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            for j, cell_text in enumerate(header_cells):
                if j < num_cols:
                    cell = table.rows[0].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.bold = True
                    run.font.size = Pt(10)

            # Data rows
            for r_idx, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    if j < num_cols:
                        cell = table.rows[r_idx + 1].cells[j]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _apply_inline_formatting(p, cell_text)
                        for run in p.runs:
                            run.font.size = Pt(10)

            doc.add_paragraph()
            continue

        # List items
        list_match = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            text = list_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            _apply_inline_formatting(p, text)
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _apply_inline_formatting(p, line)
        i += 1

    doc.save(output_path)
    return output_path
# END_BLOCK_CONVERT
